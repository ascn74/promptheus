"""Tests for attachment text extraction.

Fixtures are generated in-process rather than committed, so the repository does
not carry opaque binaries that nobody can review.
"""

import io
import zipfile

from docx import Document
from pypdf import PdfReader, PdfWriter

from promptheus.attachments import Attachment, detect_kind, extract


def make_text_pdf(pages: list[str]) -> bytes:
    """Build a minimal PDF that really does carry a text layer.

    Written by hand because pypdf can create pages but not text content, and
    pulling in a full PDF generator for four fixtures is not worth it.
    """
    page_count = len(pages)
    font_id = 3 + 2 * page_count
    kids = " ".join(f"{3 + 2 * index} 0 R" for index in range(page_count))

    objects: list[bytes] = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        f"<< /Type /Pages /Kids [{kids}] /Count {page_count} >>".encode(),
    ]
    for index, content in enumerate(pages):
        page_id = 3 + 2 * index
        objects.append(
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 200 200] "
            f"/Contents {page_id + 1} 0 R "
            f"/Resources << /Font << /F1 {font_id} 0 R >> >> >>".encode()
        )
        stream = f"BT /F1 12 Tf 20 100 Td ({content}) Tj ET".encode()
        objects.append(
            b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream"
        )
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{number} 0 obj\n".encode() + body + b"\nendobj\n"

    xref_offset = len(out)
    size = len(objects) + 1
    out += f"xref\n0 {size}\n".encode()
    out += b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode()
    out += f"trailer\n<< /Size {size} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode()
    return bytes(out)


def make_scanned_pdf(page_count: int = 2) -> bytes:
    """A PDF with pages but no text layer — what a scanner produces."""
    writer = PdfWriter()
    for _ in range(page_count):
        writer.add_blank_page(width=200, height=200)
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def make_encrypted_pdf(password: str = "secret") -> bytes:
    writer = PdfWriter()
    writer.append(PdfReader(io.BytesIO(make_text_pdf(["Confidential"]))))
    writer.encrypt(password)
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def make_docx(paragraphs: list[str], table: list[list[str]] | None = None) -> bytes:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table:
        docx_table = document.add_table(rows=len(table), cols=len(table[0]))
        for row_index, row in enumerate(table):
            for column_index, value in enumerate(row):
                docx_table.cell(row_index, column_index).text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --- kind detection ----------------------------------------------------------


def test_kind_comes_from_the_extension() -> None:
    assert detect_kind("report.pdf", b"") == "pdf"
    assert detect_kind("REPORT.PDF", b"") == "pdf"
    assert detect_kind("notes.docx", b"") == "docx"
    assert detect_kind("main.py", b"print()") == "text"


def test_magic_bytes_win_when_the_extension_lies() -> None:
    # Browsers report whatever the filesystem says; the bytes are the truth.
    assert detect_kind("mislabelled.txt", make_text_pdf(["Hi"])) == "pdf"
    assert detect_kind("mislabelled.txt", make_docx(["Hi"])) == "docx"


def test_unknown_extensions_are_treated_as_text() -> None:
    assert detect_kind("Makefile", b"all:\n\techo hi") == "text"
    assert detect_kind("weird.xyzzy", b"plain content") == "text"


def test_a_plain_zip_is_not_mistaken_for_docx() -> None:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("readme.txt", "not a word document")

    assert detect_kind("archive.zip", buffer.getvalue()) == "text"


# --- text --------------------------------------------------------------------


def test_utf8_text_round_trips() -> None:
    result = extract("notes.md", "# Título\n\nAcentuação e emoji 🔥".encode())

    assert result.kind == "text"
    assert "Título" in result.text
    assert "🔥" in result.text
    assert result.warning is None


def test_utf8_bom_is_stripped() -> None:
    result = extract("bom.txt", b"\xef\xbb\xbfhello")

    assert result.text == "hello"


def test_utf16_is_decoded_rather_than_mangled() -> None:
    result = extract("utf16.txt", "olá mundo".encode("utf-16"))

    # As latin-1 this would come back riddled with null bytes.
    assert result.text == "olá mundo"
    assert "\x00" not in result.text


def test_latin1_text_falls_back_with_a_warning() -> None:
    result = extract("legacy.txt", "ação".encode("latin-1"))

    assert result.text == "ação"
    assert result.warning is not None
    assert "latin-1" in result.warning


def test_undecodable_bytes_never_raise() -> None:
    result = extract("blob.bin", bytes(range(256)))

    assert isinstance(result, Attachment)
    assert result.kind == "text"


# --- PDF ---------------------------------------------------------------------


def test_pdf_with_a_text_layer_is_extracted() -> None:
    result = extract("doc.pdf", make_text_pdf(["Hello from page one"]))

    assert result.kind == "pdf"
    assert "Hello from page one" in result.text
    assert result.warning is None


def test_pdf_pages_are_joined() -> None:
    result = extract("doc.pdf", make_text_pdf(["First page", "Second page"]))

    assert "First page" in result.text
    assert "Second page" in result.text
    assert result.text.index("First page") < result.text.index("Second page")


def test_scanned_pdf_warns_instead_of_silently_sending_nothing() -> None:
    result = extract("scan.pdf", make_scanned_pdf(page_count=3))

    assert result.is_empty
    assert result.warning is not None
    assert "scanned" in result.warning
    assert "3 page(s)" in result.warning


def test_pdf_line_padding_is_stripped() -> None:
    # Print pipelines pad every line to a fixed column width. Those spaces
    # tokenise, and the cost is paid once per model in the run.
    padded = "Real content" + " " * 200
    result = extract("padded.pdf", make_text_pdf([padded]))

    assert "Real content" in result.text
    assert "  " not in result.text
    assert not any(line != line.rstrip() for line in result.text.splitlines())


def test_pdf_blank_line_runs_collapse() -> None:
    result = extract("gaps.pdf", make_text_pdf(["Top", "Bottom"]))

    assert "\n\n\n" not in result.text


def test_encrypted_pdf_is_reported_not_raised() -> None:
    result = extract("locked.pdf", make_encrypted_pdf())

    assert result.text == ""
    assert result.warning is not None
    assert "password-protected" in result.warning


def test_corrupt_pdf_is_reported_not_raised() -> None:
    result = extract("broken.pdf", b"%PDF-1.4\nthis is not a real pdf")

    assert result.kind == "pdf"
    assert result.text == ""
    assert result.warning is not None


# --- DOCX --------------------------------------------------------------------


def test_docx_paragraphs_are_extracted() -> None:
    result = extract("notes.docx", make_docx(["First para", "Second para"]))

    assert result.kind == "docx"
    assert "First para" in result.text
    assert "Second para" in result.text
    assert result.warning is None


def test_docx_table_content_is_not_lost() -> None:
    data = make_docx(
        ["Intro paragraph"],
        table=[["Model", "Price"], ["Opus 5", "$5"]],
    )

    result = extract("table.docx", data)

    # `document.paragraphs` skips tables entirely — the easiest way to lose
    # half a document without noticing.
    assert "Intro paragraph" in result.text
    assert "Model | Price" in result.text
    assert "Opus 5 | $5" in result.text


def test_docx_keeps_document_order() -> None:
    data = make_docx(["Before the table"], table=[["cell"]])
    result = extract("ordered.docx", data)

    assert result.text.index("Before the table") < result.text.index("cell")


def test_corrupt_docx_is_reported_not_raised() -> None:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("word/document.xml", "<not-really-a-document/>")

    result = extract("broken.docx", buffer.getvalue())

    assert result.text == ""
    assert result.warning is not None


# --- truncation --------------------------------------------------------------


def test_text_past_the_cap_is_truncated_and_flagged() -> None:
    result = extract("big.txt", b"x" * 5000, max_chars=1000)

    assert len(result.text) == 1000
    assert result.warning is not None
    assert "truncated" in result.warning


def test_text_within_the_cap_is_untouched() -> None:
    result = extract("small.txt", b"x" * 999, max_chars=1000)

    assert len(result.text) == 999
    assert result.warning is None


def test_truncation_warning_combines_with_others() -> None:
    result = extract("legacy.txt", "ação".encode("latin-1") * 500, max_chars=100)

    assert result.warning is not None
    assert "latin-1" in result.warning
    assert "truncated" in result.warning
